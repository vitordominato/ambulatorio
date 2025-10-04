# services/medicos_parser.py
from io import BytesIO
from typing import List, Dict
import re
import yaml
import pandas as pd
from docx import Document

# ------------ utilidades de texto ------------
def _norm(s: str) -> str:
    return (s or "").strip()

def _is_dias_line(t: str) -> bool:
    t = t.lower()
    return "dia" in t and ("atendimento" in t or "atende" in t)

def _parse_dias(t: str) -> list:
    # aceita "Dias de atendimento: Segunda e Quinta / Terça, Sexta"
    if ":" in t:
        t = t.split(":", 1)[1]
    parts = re.split(r"[;,/]| e ", t, flags=re.IGNORECASE)
    return [p.strip().capitalize() for p in parts if p.strip()]

def _looks_like_spec(t: str) -> bool:
    """Heurística ampla p/ especialidade: curta, sem dois-pontos, 1-4 palavras, cada uma inicial maiúscula."""
    if ":" in t or len(t) > 60: 
        return False
    words = t.split()
    if not (1 <= len(words) <= 6):
        return False
    # maioria das palavras começa com maiúscula (Cardiologia, Cirurgia Geral etc.)
    caps = sum(1 for w in words if w[:1].isupper())
    return caps >= max(1, int(0.7 * len(words)))

# ------------ coleta linear do conteúdo (parágrafos + tabelas) ------------
def _collect_lines(doc: Document) -> List[str]:
    lines = []
    # parágrafos
    for p in doc.paragraphs:
        txt = _norm(p.text)
        if txt:
            lines.append(txt)
    # tabelas (linha a linha)
    for tbl in doc.tables:
        for row in tbl.rows:
            cell_texts = []
            seen = set()
            for cell in row.cells:
                # evita duplicatas por merges internos do python-docx
                if id(cell) in seen: 
                    continue
                seen.add(id(cell))
                t = _norm(cell.text)
                if t:
                    cell_texts.append(t)
            if cell_texts:
                # Se a linha tem "Dias..." em alguma célula, preserva cada célula separada
                if any(_is_dias_line(t) for t in cell_texts) and len(cell_texts) >= 2:
                    lines.extend(cell_texts)
                else:
                    lines.append(" | ".join(cell_texts))
    return lines

# ------------ parser principal ------------
def parse_docx_bytes(docx_bytes: bytes) -> List[Dict]:
    doc = Document(BytesIO(docx_bytes))
    lines = _collect_lines(doc)

    records: List[Dict] = []
    current_spec = None
    last_nonempty = None
    last_name = None

    for raw in lines:
        t = _norm(raw)

        # especialidade (heurística)
        if _looks_like_spec(t) and not _is_dias_line(t):
            current_spec = t
            last_name = None
            last_nonempty = t
            continue

        # linha com "Dias..."
        if _is_dias_line(t):
            dias = _parse_dias(t)
            # tenta achar o nome na célula anterior/mesma linha
            name_candidate = last_name or (last_nonempty if last_nonempty and last_nonempty != current_spec else None)
            if " | " in raw:  # linha veio de tabela com duas células (nome | dias)
                left = raw.split(" | ", 1)[0].strip()
                if left and not _is_dias_line(left):
                    name_candidate = left
            if current_spec and name_candidate:
                records.append({"especialidade": current_spec, "medico": name_candidate, "dias": dias})
                last_name = None
                last_nonempty = t
                continue

        # possível nome de médico
        if t and not _is_dias_line(t):
            # evita confundir especialidade como nome
            if current_spec and t != current_spec:
                last_name = t
        last_nonempty = t

    # sobrou um nome sem dias → registra assim mesmo
    if current_spec and last_name:
        records.append({"especialidade": current_spec, "medico": last_name, "dias": []})

    return records

# ------------ exportadores ------------
def to_yaml_by_specialty(records: List[Dict]) -> str:
    out = {}
    for r in records:
        out.setdefault(r["especialidade"], []).append({"nome": r["medico"], "dias": r["dias"]})
    return yaml.safe_dump(out, allow_unicode=True, sort_keys=True)

def to_csv(records: List[Dict]) -> str:
    if not records:
        return ""
    df = pd.DataFrame(records)
    if "dias" not in df.columns:
        df["dias"] = [[] for _ in range(len(df))]
    df["dias"] = df["dias"].apply(lambda x: ", ".join(x) if isinstance(x, list) else str(x))
    return df.to_csv(index=False)
